from __future__ import annotations

from pathlib import Path
import sys

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.config import OUTPUT_DIR, PROJECT_ROOT


REPORT_PATH = PROJECT_ROOT / "report" / "VAST2023_MC1_非法捕鱼知识图谱可视分析实验报告_FishHook统一版.docx"


def east_asia(run, font_name: str = "Microsoft YaHei") -> None:
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": 80, "start": 120, "bottom": 80, "end": 120}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for name, size, color, before, after in [
        ("Heading 1", 16, "1F4D78", 14, 8),
        ("Heading 2", 13, "2E74B5", 10, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def p(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    east_asia(run)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        east_asia(run)


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        east_asia(run)


def table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for idx, header in enumerate(headers):
        cell = t.rows[0].cells[idx]
        cell.text = str(header)
        shade(cell, "E8EEF5")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Calibri"
                east_asia(run)
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = "" if pd.isna(value) else str(value)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    east_asia(run)
    for row in t.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            margins(row.cells[idx])


def clean(value: object, n: int = 150) -> str:
    text = str(value)
    return text if len(text) <= n else text[: n - 3] + "..."


def first_existing(df: pd.DataFrame, names: list[str], default: object = "") -> pd.Series:
    for name in names:
        if name in df.columns:
            return df[name]
    return pd.Series([default] * len(df))


def main() -> None:
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    nodes = pd.read_csv(OUTPUT_DIR / "nodes.csv")
    edges = pd.read_csv(OUTPUT_DIR / "edges.csv")
    suspects = pd.read_csv(OUTPUT_DIR / "suspect_summary.csv")
    expansion = pd.read_csv(OUTPUT_DIR / "suspect_expansion_candidates.csv")
    fishhook = pd.read_csv(OUTPUT_DIR / "fishhook_candidate_ranking.csv")
    common = pd.read_csv(OUTPUT_DIR / "suspect_common_neighbors.csv")
    companies = pd.read_csv(OUTPUT_DIR / "company_fishhook_ranking.csv")

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VAST Challenge 2023 MC1 非法捕鱼知识图谱可视分析实验报告")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    east_asia(run)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于 FishHook 思路的可疑实体扩展、异常评分与多视图分析")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("555555")
    east_asia(run)

    doc.add_heading("摘要", level=1)
    p(
        doc,
        "本实验选择 VAST Challenge 2023 Mini-Challenge 1 作为课程设计题目，围绕非法捕鱼调查任务构建知识图谱可视分析系统。"
        "系统参考 FishHook 论文的分析思路，以官方给出的四个可疑实体为调查种子，结合 Ego Net、候选实体扩展、FishHook 风格异常特征、"
        "平行坐标对比、证据路径图和共同邻居图，识别更值得 FishEye 继续调查的公司、船只、人物和组织。"
    )

    doc.add_heading("1. 实验目的", level=1)
    bullets(
        doc,
        [
            "理解 VAST 2023 MC1 知识图谱中的节点、边和关系类型。",
            "实现一个可交互的非法捕鱼知识图谱可视分析系统。",
            "从官方可疑实体出发，发现其他潜在非法捕鱼相关实体。",
            "参考 FishHook 论文，将异常评分、短环路、弱连通分量船只比例和平行坐标图引入课程项目。",
            "形成可用于报告和课堂展示的证据链、候选实体排序和可视化结论。",
        ],
    )

    doc.add_heading("2. 数据集与属性说明", level=1)
    p(
        doc,
        f"原始数据文件为 MC1/V1/MC1.json，采用 node-link 形式表示有向多重知识图谱。"
        f"当前项目读取到 {len(nodes)} 个节点和 {len(edges)} 条边。节点类型包括公司、人物、组织、政治组织、地点、船只、事件等；"
        "边类型包括 ownership、membership、partnership 和 family_relationship。"
    )
    node_counts = nodes["type"].fillna("unknown").value_counts().reset_index()
    node_counts.columns = ["节点类型", "数量"]
    edge_counts = edges["type"].fillna("unknown").value_counts().reset_index()
    edge_counts.columns = ["边类型", "数量"]
    table(doc, ["节点类型", "数量"], node_counts.values.tolist(), [3.25, 3.25])
    p(doc, "表 1 节点类型统计。")
    table(doc, ["边类型", "数量"], edge_counts.values.tolist(), [3.25, 3.25])
    p(doc, "表 2 边类型统计。")
    table(
        doc,
        ["关系类型", "含义", "在非法捕鱼调查中的解释"],
        [
            ["ownership", "所有权或控制关系", "用于识别公司、组织和资产背后的控制结构。"],
            ["membership", "成员或隶属关系", "用于识别组织网络和实体归属。"],
            ["partnership", "合作关系", "用于识别商业往来和合作网络。"],
            ["family_relationship", "家庭或亲属关系", "用于发现可能的隐性控制或私人关系网络。"],
        ],
        [1.45, 2.0, 3.05],
    )

    doc.add_heading("3. 系统设计与功能", level=1)
    numbered(
        doc,
        [
            "Ego Net：围绕当前实体展示一跳或二跳关系网络，并高亮官方可疑实体。",
            "FishHook 候选排名：综合多个结构特征，对候选实体进行调查优先级排序。",
            "平行坐标对比：同时比较多个候选实体在多种异常特征上的表现。",
            "从可疑实体扩展：以四个官方可疑实体为种子，寻找 1-3 跳内的候选对象。",
            "证据路径图：用 Sankey 图展示当前实体到非法捕鱼关键词相关节点的短路径。",
            "共同邻居分析：用热力图和二部网络图展示可疑实体之间共享的中介节点。",
        ],
    )

    doc.add_heading("4. FishHook 风格异常特征", level=1)
    p(
        doc,
        "FishHook 论文强调通过结构特征和异常模式追踪可疑实体。本项目在课程实现中保留了其核心思想，"
        "但采用本地可直接计算的图指标近似实现，包括弱连通分量船只比例、1-2 跳家族关系、1-2 跳政治组织、短环路数量、"
        "与官方可疑实体的距离、船只连接、所有权边和成员边等。"
    )
    table(
        doc,
        ["指标", "计算含义", "解释"],
        [
            ["异常分", "综合异常得分", "用于整体排序，分数越高越值得优先调查。"],
            ["船只比例", "所在弱连通分量中 vessel 节点占比", "反映候选实体所在团簇是否偏渔业供应链。"],
            ["家族边", "1-2 跳范围内 family_relationship 边数量", "提示是否存在家族或私人关系网络。"],
            ["政治组织", "1-2 跳范围内政治组织节点数量", "提示与政治组织或权力结构的接近程度。"],
            ["短环路", "经过该实体的短有向环路数量", "提示复杂控制结构、循环供应链或异常闭环。"],
            ["关联嫌疑", "连接到的官方可疑实体数量", "说明该候选实体与已知嫌疑对象的重叠程度。"],
            ["船只连接", "直接相连的 vessel 节点数量", "越高越接近非法捕鱼调查场景。"],
            ["所有权边", "直接相关 ownership 边数量", "反映控制或所有权结构。"],
            ["成员边", "直接相关 membership 边数量", "反映组织归属关系。"],
        ],
        [1.1, 2.2, 3.2],
    )

    doc.add_heading("5. 平行坐标图解释", level=1)
    p(
        doc,
        "平行坐标图中的每一根竖轴代表一个特征指标，每一条折线代表一个候选实体。"
        "如果某条折线在多个关键轴上都处于较高位置，说明该实体并非单一指标异常，而是在多个结构维度上都具有较强可疑性。"
        "因此，平行坐标图适合用于比较候选实体的异常模式和筛选优先调查对象。"
    )
    p(
        doc,
        "需要注意的是，平行坐标图不直接证明违法行为。它的作用是帮助分析人员发现“多指标同时偏高”的对象，"
        "再结合 Ego Net、证据路径、共同邻居和候选排名进一步解释证据链。"
    )

    doc.add_heading("6. 实验结果", level=1)
    doc.add_heading("6.1 官方可疑实体分析", level=2)
    suspect_cols = [
        "entity",
        "fishhook_anomaly_score",
        "component_vessel_ratio",
        "family_edges_1_2_hop",
        "political_nodes_1_2_hop",
        "short_cycle_count",
        "vessel_links",
        "ownership_edges",
        "membership_edges",
    ]
    table(
        doc,
        ["实体", "FishHook分", "船只比例", "家族边", "政治组织", "短环路", "船只连接", "所有权边", "成员边"],
        suspects[suspect_cols].values.tolist(),
        [1.25, 0.65, 0.65, 0.55, 0.6, 0.55, 0.55, 0.55, 0.55],
    )
    p(doc, "官方可疑实体统一展示 FishHook 风格异常特征，用于说明其结构异常来源。")

    doc.add_heading("6.2 FishHook 候选排名", level=2)
    top = fishhook.head(10).copy()
    top["path_examples"] = top["path_examples"].map(lambda x: clean(x, 120))
    top_display = pd.DataFrame(
        {
            "candidate": top["candidate"],
            "type": top["type"],
            "fishhook_anomaly_score": top["fishhook_anomaly_score"],
            "connected_suspect_count": first_existing(top, ["connected_suspect_count", "connected_suspect_count_x"]),
            "short_cycle_count": top["short_cycle_count"],
            "vessel_links": first_existing(top, ["vessel_links", "vessel_links_x"]),
            "path_examples": top["path_examples"],
        }
    )
    table(
        doc,
        ["候选实体", "类型", "FishHook分", "关联嫌疑", "短环路", "船只连接", "路径示例"],
        top_display.values.tolist(),
        [1.35, 0.65, 0.75, 0.65, 0.55, 0.55, 1.95],
    )
    p(
        doc,
        "排名靠前的对象通常同时满足多个条件：接近多个官方可疑实体、存在较多船只连接、短环路数量较多，"
        "并且具有 ownership 或 membership 等结构性关系。"
    )

    doc.add_heading("6.3 共同邻居分析", level=2)
    table(
        doc,
        ["实体 A", "实体 B", "共同邻居数", "共同邻居类型", "示例"],
        common.head(8).values.tolist(),
        [1.25, 1.35, 0.75, 1.3, 1.85],
    )
    p(doc, "共同邻居越多，说明两个可疑实体越可能共享同一中介、地点、组织或关系团簇。")

    doc.add_heading("6.4 推荐调查公司", level=2)
    company_display = pd.DataFrame(
        {
            "candidate": companies["candidate"],
            "fishhook_anomaly_score": companies["fishhook_anomaly_score"],
            "connected_suspect_count": first_existing(companies, ["connected_suspect_count", "connected_suspect_count_x"]),
            "short_cycle_count": companies["short_cycle_count"],
            "vessel_links": first_existing(companies, ["vessel_links", "vessel_links_x"]),
        }
    ).head(8)
    table(
        doc,
        ["公司", "FishHook分", "关联嫌疑", "短环路", "船只连接"],
        company_display.values.tolist(),
        [1.9, 0.8, 0.8, 0.75, 0.75],
    )

    doc.add_heading("7. 结论", level=1)
    bullets(
        doc,
        [
            "从官方可疑实体出发进行扩展，比单纯关键词匹配更符合 MC1 的调查逻辑。",
            "FishHook 风格异常特征能够帮助识别多指标同时异常的候选对象。",
            "平行坐标图适合比较候选实体的异常模式，证据路径图适合解释语义线索，共同邻居图适合发现共享中介。",
            "当前结果只能表示调查优先级，不能直接证明违法；最终结论仍需要结合原始文本语境进一步验证。",
        ],
    )

    doc.add_heading("8. 不足与改进", level=1)
    bullets(
        doc,
        [
            "原始知识图谱由 NLP 自动抽取得到，存在乱码、匿名 ID 和关系噪声。",
            "短环路检测采用有界搜索以保证交互速度，不能穷举所有环路。",
            "关键词路径属于辅助线索，不能作为直接证据。",
            "后续可加入社区发现、图嵌入距离、路径高亮和原文证据关联。",
        ],
    )

    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    main()
